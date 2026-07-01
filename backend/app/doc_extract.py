# =====================================================================
#  HUJJATDAN MATN AJRATISH (PDF / DOCX / TXT)
#  Skaner qilingan (rasm) PDF'lar uchun matn chiqmasligi mumkin — OCR
#  hozircha qo'llab-quvvatlanmaydi.
# =====================================================================
import io


class ExtractError(Exception):
    pass


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        try:
            import pdfplumber
        except ImportError:
            raise ExtractError("pdfplumber o'rnatilmagan (requirements.txt)")
        out = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                out.append(page.extract_text() or "")
        return "\n".join(out).strip()

    if name.endswith(".docx"):
        try:
            import docx  # python-docx
        except ImportError:
            raise ExtractError("python-docx o'rnatilmagan (requirements.txt)")
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs]
        # jadvallardagi matn ham
        for tbl in d.tables:
            for row in tbl.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts).strip()

    if name.endswith(".txt") or name.endswith(".md"):
        return data.decode("utf-8", "ignore").strip()

    raise ExtractError("Qo'llab-quvvatlanmaydigan fayl turi. Faqat: PDF, DOCX, TXT.")
